from app.schemas.document import Document
from app.schemas.document_chunk import DocumentChunk
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

def read_text_file(file_path:str) -> str:
    path = Path(file_path)
    return path.read_text(encoding="utf-8")

def _read_pdf_file(file_path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p.strip() for p in pages if p.strip())

def _read_docx_file(file_path: str) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)

def read_document_file(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix in {".txt", ".md"}:
        return read_text_file(file_path)
    if suffix == ".pdf":
        return _read_pdf_file(file_path)
    if suffix == ".docx":
        return _read_docx_file(file_path)
    raise ValueError(f"Unsupported file type: {suffix}")

def create_document(
    document_id:int,
    filename:str,
    file_type:str,
    content:str,
) -> Document:
    document = Document(
        id = document_id,
        filename = filename,
        file_type = file_type,
        content = content,
    )
    return document

def create_document_from_file(
    document_id:int,
    file_path:str,
    ) -> Document:
    path = Path(file_path)
    content = read_document_file(file_path)

    return create_document(
        document_id=document_id,
        filename=path.name,
        file_type=path.suffix.lstrip("."),
        content=content,
        )

def preview_document(document:Document,max_length:int = 30) -> str:
    if len(document.content) <= max_length:
        return document.content
    
    return document.content[:max_length] + "..."

def create_document_chunk(
    chunk_id:int,
    document_id:int,
    chunk_index:int,
    content:str,
) -> DocumentChunk:
    chunk = DocumentChunk(
        id = chunk_id,
        document_id = document_id,
        chunk_index = chunk_index,
        content = content,
    )
    return chunk

def split_document_into_chunks(
    document:Document,
    chunk_size:int = 20,
    overlap : int = 0,
) -> list[DocumentChunk]:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be greater than 0")
    if overlap < 0:
        raise ValueError("overlap must be greater than or equal to 0")  
    if overlap >= chunk_size:
        raise ValueError("overlap must be smaller than chunk_size")  
    
    if not document.content:
        return[]
    
    chunks = []
    step = chunk_size - overlap

    for start in range(0,len(document.content),step):
        end = start+chunk_size
        chunk_content = document.content[start:end]
        chunk_index = len(chunks)

        chunk = create_document_chunk(
            chunk_id=chunk_index+1,
            document_id=document.id,
            chunk_index=chunk_index,
            content=chunk_content,
            )

        chunks.append(chunk)

    return chunks

def _atomize(text: str, max_len: int) -> list[str]:
    """把一段文本拆成不超过 max_len 的原子片段：先按换行，再按句末标点，最后按长度硬切。"""
    text = text.strip()
    if not text:
        return []
    if len(text) <= max_len:
        return [text]

    # 逐级拆分：先换行分行，超长的行再按句末标点分句
    units: list[str] = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if len(line) <= max_len:
            units.append(line)
            continue
        sentence = ""
        for ch in line:
            sentence += ch
            if ch in "。！？!?；;" and len(sentence) >= max_len // 2:
                units.append(sentence.strip())
                sentence = ""
        if sentence.strip():
            units.append(sentence.strip())

    # 仍有超长单元（长句无标点）则按长度硬切
    result: list[str] = []
    for u in units:
        if len(u) <= max_len:
            result.append(u)
        else:
            for i in range(0, len(u), max_len):
                result.append(u[i:i + max_len])
    return result


def split_document_by_paragraphs(
    document: Document,
    min_chunk_len: int = 60,
    max_chunk_len: int = 250,
) -> list[DocumentChunk]:
    raw_paragraphs = [p.strip() for p in document.content.split("\n\n") if p.strip()]
    if not raw_paragraphs:
        return []

    # 先把每个段落拆成不超过 max_chunk_len 的原子片段，
    # 避免 PDF 整页文本变成一个巨大杂糅块、稀释语义导致检索不到。
    atoms: list[str] = []
    for para in raw_paragraphs:
        atoms.extend(_atomize(para, max_chunk_len))

    # 再把过短的原子片段（标题、目录行等）累积合并到 min_chunk_len，
    # 但合并后不超过 max_chunk_len，兼顾"不碎"与"不过大"。
    merged: list[str] = []
    buffer = ""
    for atom in atoms:
        candidate = f"{buffer}\n{atom}" if buffer else atom
        if len(candidate) > max_chunk_len and buffer:
            merged.append(buffer)
            buffer = atom
        else:
            buffer = candidate
        if len(buffer) >= min_chunk_len:
            merged.append(buffer)
            buffer = ""
    if buffer:
        if merged and len(buffer) < min_chunk_len:
            merged[-1] = f"{merged[-1]}\n{buffer}"
        else:
            merged.append(buffer)

    chunks = []
    for content in merged:
        chunk_index = len(chunks)
        chunks.append(
            create_document_chunk(
                chunk_id=chunk_index + 1,
                document_id=document.id,
                chunk_index=chunk_index,
                content=content,
            )
        )
    return chunks

def list_text_files(directory_path:str) -> list[str]:
    directory = Path(directory_path)
    if not directory.exists():
        return[]

    text_files = []
    for file_path in sorted(directory.iterdir()):
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            text_files.append(str(file_path))

    return text_files


def _sanitize_path_segment(name: str, fallback: str = "unnamed") -> str:
    """把用户名/知识库名清洗成一个安全的单层目录名。

    - 替换文件系统非法字符（\\ / : * ? " < > |）与控制字符为 _
    - 去掉首尾空格与点（Windows 目录名不能以点/空格结尾）
    - 防路径穿越：结果绝不含 / \\ 或 ..（杜绝 ../ 逃逸）
    - 清洗后为空则用 fallback
    """
    import re

    name = str(name or "")
    # 非法字符与控制字符 → _
    name = re.sub(r'[\\/:*?"<>|\x00-\x1f]', "_", name)
    # .. 是穿越关键，直接抹掉
    name = name.replace("..", "_")
    # 去首尾空格和点
    name = name.strip().strip(".").strip()
    if not name:
        return fallback
    # 限制长度，避免超出文件系统上限
    return name[:80]


def kb_documents_dir(kb_id: int) -> Path:
    """某个知识库的文件目录：DOCUMENTS_DIR/{用户名}/{知识库名}_{kb_id}/。

    多知识库隔离后，每个库的文件物理分目录存放。目录名用「用户名/库名_kbid」
    便于人工在文件系统里辨认；末尾的 _kbid 保证同名库不冲突、天然唯一。
    需要时自动创建目录。

    降级兜底：查不到知识库或 DB 不可用时，回退到 DOCUMENTS_DIR/{kb_id}/，
    保证上传/问答主流程不因目录解析失败而中断。
    """
    from app.config import DOCUMENTS_DIR

    base = Path(DOCUMENTS_DIR)
    try:
        # 局部导入避免与 kb_service/user_service 的循环依赖
        from app.services import kb_service, user_service

        kb = kb_service.get(kb_id)
        if kb:
            owner = user_service.get_by_id(kb["owner_id"])
            username = owner["username"] if owner else f"user{kb['owner_id']}"
            directory = (
                base
                / _sanitize_path_segment(username, fallback=f"user{kb['owner_id']}")
                / f"{_sanitize_path_segment(kb['name'], fallback='kb')}_{kb_id}"
            )
            directory.mkdir(parents=True, exist_ok=True)
            return directory
    except Exception:
        # 任何异常都走降级，绝不中断主流程
        pass

    directory = base / str(kb_id)
    directory.mkdir(parents=True, exist_ok=True)
    return directory